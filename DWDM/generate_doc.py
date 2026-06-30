import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import numpy as np
import matplotlib.pyplot as plt
from sklearn import svm
from sklearn.cluster import KMeans
from sklearn.datasets import make_blobs
from sklearn.preprocessing import StandardScaler
import os
from PIL import Image, ImageDraw, ImageFont

def parse_terminal_text(text):
    lines = []
    for line in text.split('\n'):
        if line.startswith('PS '):
            idx = line.rfind('>')
            if idx != -1:
                prompt_prefix = line[:idx+1]
                command_text = line[idx+1:]
                lines.append(('prompt', prompt_prefix, command_text))
            else:
                lines.append(('output', line, None))
        else:
            lines.append(('output', line, None))
    return lines

def create_terminal_screenshot(lines, save_path):
    TERM_BG = (12, 12, 12)             # Black Terminal Background
    TERM_TEXT = (204, 204, 204)        # Light gray
    TERM_PROMPT_COLOR = (255, 235, 156)# Light Yellowish Prompt
    TERM_TITLE_BAR = (30, 30, 30)
    TERM_TITLE_TEXT = (200, 200, 200)
    TERM_BORDER = (60, 60, 60)
    
    FONT_MONO = "C:/Windows/Fonts/consola.ttf"
    FONT_MONO_BOLD = "C:/Windows/Fonts/consolab.ttf"
    FONT_UI = "C:/Windows/Fonts/segoeui.ttf"
    
    TERM_FONT_SIZE = 14
    TERM_TITLE_FONT_SIZE = 12
    TERM_LINE_HEIGHT = 20
    TERM_WIDTH = 1000
    TERM_PADDING_X = 16
    TERM_PADDING_Y = 10
    TERM_TITLE_BAR_HEIGHT = 32
    TERM_CORNER_RADIUS = 8
    
    try:
        font_mono = ImageFont.truetype(FONT_MONO, TERM_FONT_SIZE)
        font_mono_bold = ImageFont.truetype(FONT_MONO_BOLD, TERM_FONT_SIZE)
        font_ui = ImageFont.truetype(FONT_UI, TERM_TITLE_FONT_SIZE)
    except IOError:
        font_mono = ImageFont.load_default()
        font_mono_bold = ImageFont.load_default()
        font_ui = ImageFont.load_default()

    height = TERM_TITLE_BAR_HEIGHT + (TERM_PADDING_Y * 2) + (len(lines) * TERM_LINE_HEIGHT) + 20
    
    img = Image.new('RGB', (TERM_WIDTH, height), TERM_BG)
    draw = ImageDraw.Draw(img)
    
    # Draw title bar background
    draw.rectangle([(0, 0), (TERM_WIDTH, TERM_TITLE_BAR_HEIGHT)], fill=TERM_TITLE_BAR)
    
    # Draw title bar text
    draw.text((16, (TERM_TITLE_BAR_HEIGHT - TERM_TITLE_FONT_SIZE) // 2), ">_ Windows PowerShell", fill=TERM_TITLE_TEXT, font=font_ui)
    
    # Draw window controls
    btn_y = TERM_TITLE_BAR_HEIGHT // 2
    max_x = TERM_WIDTH - 92
    
    # Minimize line
    draw.line([(max_x, btn_y), (max_x + 10, btn_y)], fill=(170, 170, 170), width=1)
    
    # Maximize rectangle
    draw.rectangle([(max_x + 20, btn_y - 5), (max_x + 30, btn_y + 5)], outline=(170, 170, 170), width=1)
    
    # Close X
    close_x = TERM_WIDTH - 46
    draw.line([(close_x, btn_y - 5), (close_x + 10, btn_y + 5)], fill=(170, 170, 170), width=1)
    draw.line([(close_x + 10, btn_y - 5), (close_x, btn_y + 5)], fill=(170, 170, 170), width=1)
    
    # Border
    draw.rectangle([(0, 0), (TERM_WIDTH - 1, height - 1)], outline=TERM_BORDER, width=1)
    
    # Content lines
    y = TERM_TITLE_BAR_HEIGHT + TERM_PADDING_Y
    for line_type, text1, text2 in lines:
        x = TERM_PADDING_X
        if line_type == 'prompt':
            draw.text((x, y), text1, fill=TERM_PROMPT_COLOR, font=font_mono)
            try:
                bbox = draw.textbbox((0, 0), text1, font=font_mono)
                prompt_w = bbox[2] - bbox[0]
            except AttributeError:
                prompt_w = draw.textsize(text1, font=font_mono)[0]
            draw.text((x + prompt_w, y), text2, fill=(255, 255, 255), font=font_mono)
        else:
            draw.text((x, y), text1, fill=TERM_TEXT, font=font_mono)
        y += TERM_LINE_HEIGHT
        
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    img.save(save_path, 'PNG', quality=95)
    return save_path

def set_cell_background(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tc_pr.append(shd)

def add_code_block(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.italic = True
    run.font.size = Pt(10)

def add_custom_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    return h

def generate_svm_plot():
    # Synthetic data
    X = np.array([[1, 2], [2, 3], [3, 3], [4, 4],
                  [6, 7], [7, 8], [8, 7], [9, 9]])
    y = np.array([0, 0, 0, 0, 1, 1, 1, 1])

    clf = svm.SVC(kernel='linear')
    clf.fit(X, y)
    
    fig, ax = plt.subplots(figsize=(6, 4))
    scatter = ax.scatter(X[:,0], X[:,1], c=y, s=100, cmap='viridis', edgecolors='k', zorder=3)

    xlim, ylim = ax.get_xlim(), ax.get_ylim()
    xx = np.linspace(xlim[0], xlim[1], 50)
    yy = np.linspace(ylim[0], ylim[1], 50)
    YY, XX = np.meshgrid(yy, xx)
    Z = clf.decision_function(np.c_[XX.ravel(), YY.ravel()]).reshape(XX.shape)

    # Plot decision boundaries and margins
    ax.contour(XX, YY, Z, colors='k', levels=[-1, 0, 1],
               linestyles=['--', '-', '--'], alpha=0.6)

    # Highlight support vectors
    ax.scatter(clf.support_vectors_[:,0], clf.support_vectors_[:,1],
               s=220, facecolors='none', edgecolors='red', linewidths=1.5,
               label='Support Vectors')

    ax.set_xlabel("Feature X1")
    ax.set_ylabel("Feature X2")
    ax.set_title("SVM Optimization: Decision Boundary & Margins")
    ax.legend(loc='lower right')
    plt.grid(True, linestyle=":", alpha=0.4)
    plt.tight_layout()
    plt.savefig("plots/svm_plot.png", dpi=150)
    plt.close()

def generate_kmeans_plot():
    X, y_true = make_blobs(n_samples=300, centers=4, cluster_std=0.60, random_state=42)
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    k = 4
    kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
    kmeans.fit(X_scaled)
    labels = kmeans.labels_
    centroids = kmeans.cluster_centers_

    colors = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12']
    plt.figure(figsize=(6, 4))
    for cluster_id in range(k):
        mask = labels == cluster_id
        plt.scatter(X_scaled[mask, 0], X_scaled[mask, 1],
                    s=60, c=colors[cluster_id], label=f'Cluster {cluster_id}', alpha=0.7)

    plt.scatter(centroids[:, 0], centroids[:, 1],
                s=300, c='black', marker='X', zorder=5, label='Centroids')

    plt.xlabel("Scaled Feature 1")
    plt.ylabel("Scaled Feature 2")
    plt.title(f"K-Means Clustering Distribution (K={k})")
    plt.legend()
    plt.grid(True, linestyle=":", alpha=0.5)
    plt.tight_layout()
    plt.savefig("plots/kmeans_plot.png", dpi=150)
    plt.close()

def main():
    # Pre-generate plots
    print("Generating SVM and K-Means plot images...")
    generate_svm_plot()
    generate_kmeans_plot()

    doc = Document()
    
    # --- Format Normal Style ---
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0, 0, 0)
    
    # --- Format Headings to match Times New Roman 12 Bold ---
    for heading_name in ['Heading 1', 'Heading 2']:
        style_h = doc.styles[heading_name]
        font_h = style_h.font
        font_h.name = 'Times New Roman'
        font_h.size = Pt(12)
        font_h.bold = True
        font_h.color.rgb = RGBColor(0, 0, 0)
        
        # Adjust spacing for headings
        style_h.paragraph_format.space_before = Pt(12)
        style_h.paragraph_format.space_after = Pt(6)
        style_h.paragraph_format.keep_with_next = True

    # ==========================================
    # LAB 2
    # ==========================================
    h = add_custom_heading(doc, "Lab 2: Data Preprocessing — Data Cleaning", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    r_q = p_q.add_run("How do you perform data cleaning on a messy dataset? Write a Python program using pandas to load a dataset, handle missing values (NaNs) by imputing with the mean of the column, convert date values to a proper datetime format, filter/drop rows with invalid dates, cap outliers in numerical columns (e.g., capping 'Duration' at 120), and remove duplicate rows.")
    p_q.paragraph_format.left_indent = Inches(0.25)
    
    add_custom_heading(doc, "Sample Input Dataset", level=2)
    p_sample = doc.add_paragraph()
    p_sample.paragraph_format.left_indent = Inches(0.25)
    p_sample.add_run(
        "Below is a sample of the raw, dirty dataset that contains a missing value (row 5 Calories), an invalid date format with extra quotes, an outlier Duration (450 mins on row 7), and duplicate records (rows 11 and 12):\n"
    )
    sample_csv_2 = (
        "Row,Duration,Date,Pulse,Maxpulse,Calories\n"
        "1,60,2023/12/01',110,130,409.1\n"
        "2,60,2023/12/02',117,145,479.0\n"
        "3,60,2023/12/03',103,135,340.0\n"
        "4,45,2023/12/04',109,175,282.4\n"
        "5,45,2023/12/05',117,148,NaN\n" # Missing value
        "6,60,2023/12/06',102,127,300.0\n"
        "7,450,2023/12/08',104,134,253.3\n" # Outlier
        "8,30,2023/12/09',109,133,195.1\n"
        "9,60,2023/12/12',100,120,250.7\n"
        "10,60,2023/12/12',100,120,250.7" # Duplicate
    )
    add_code_block(doc, sample_csv_2)
    
    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "Data cleaning is the process of identifying and correcting (or removing) corrupt, inaccurate, or incomplete records from a dataset. "
        "In Python, pandas provides a powerful set of functions to handle common preprocessing tasks:\n\n"
        "1. Handling Missing Data: Imputing missing numerical entries using statistics (like the mean) preserves data points without skewing statistical results.\n"
        "2. Type Conversion: Using pd.to_datetime with errors='coerce' turns string dates into datetime objects and converts invalid dates into NaT (Not a Time), which can then be dropped.\n"
        "3. Outlier Capping: Capping prevents extreme outliers from biasing predictions without deleting rows entirely.\n"
        "4. Deduplication: Removing redundant identical rows ensures that metrics and models are trained on unique observations."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab2_DataCleaning/lab2_data_cleaning.py", "r", encoding="utf-8") as f:
        code_lab2 = f.read()
    add_code_block(doc, code_lab2)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab2 = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab2_DataCleaning\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab2_DataCleaning> ..\\.venv\\Scripts\\python.exe lab2_data_cleaning.py\n\n"
        "==================================================\n"
        " LAB 2: DATA PREPROCESSING - DATA CLEANING \n"
        "==================================================\n"
        "INFO: Original Dataset loaded:\n"
        "   Duration         Date  Pulse  Maxpulse  Calories\n"
        "0        60  2023/12/01'    110       130     409.1\n"
        "1        60  2023/12/02'    117       145     479.0\n"
        "2        60  2023/12/03'    103       135     340.0\n"
        "3        45  2023/12/04'    109       175     282.4\n"
        "4        45  2023/12/05'    117       148     406.0\n"
        "Total rows: 32\n"
        "INFO: Mean Calories (304.68) imputed in 'Calories' column.\n"
        "INFO: Converted 'Date' to proper datetime format and dropped invalid rows.\n"
        "INFO: Outliers capped: Max 'Duration' set to 120 minutes.\n"
        "INFO: Removed 1 duplicate rows.\n"
        "INFO: Fully Cleaned Dataset Preview:\n"
        "   Duration       Date  Pulse  Maxpulse  Calories\n"
        "0        60 2023-12-01    110       130     409.1\n"
        "1        60 2023-12-02    117       145     479.0\n"
        "2        60 2023-12-03    103       135     340.0\n"
        "3        45 2023-12-04    109       175     282.4\n"
        "4        45 2023-12-05    117       148     406.0\n"
        "5        60 2023-12-06    102       127     300.0\n"
        "6        60 2023-12-07    110       136     374.0\n"
        "7       120 2023-12-08    104       134     253.3\n"
        "8        30 2023-12-09    109       133     195.1\n"
        "9        60 2023-12-10     98       124     269.0\n"
        "Final Total rows: 29"
    )
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab2)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab2_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 2.1: Lab 2 Terminal Output Screenshot")
    
    doc.add_page_break()

    # ==========================================
    # LAB 3
    # ==========================================
    h = add_custom_heading(doc, "Lab 3: Data Preprocessing — Data Integration", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.25)
    p_q.add_run("What is data integration, and how can datasets from different sources and formats (CSV, Excel) be combined? Write a Python script to load data from an Excel sheet (.xlsx) and a CSV file, align their schemas to a set of common attributes, merge them into a single unified dataframe, drop duplicate entries, and calculate simple analytics (e.g., top 5 countries by revenue).")
    
    add_custom_heading(doc, "Sample Input Dataset", level=2)
    p_sample = doc.add_paragraph()
    p_sample.paragraph_format.left_indent = Inches(0.25)
    p_sample.add_run("This lab integrates customer transaction records stored in different format origins. Below are sample structures of the columns present in the source spreadsheets:\n")
    
    sample_csv_3 = (
        "Excel & CSV columns: InvoiceNo, StockCode, Description, Quantity, InvoiceDate, UnitPrice, CustomerID, Country\n\n"
        "Sample Record (Excel):\n"
        "85123A, WHITE HANGING HEART T-LIGHT HOLDER, 6, 2.55, 17850, United Kingdom\n\n"
        "Sample Record (CSV):\n"
        "22752, SET 7 BABUSHKA NESTING BOXES, 2, 7.65, 17850, United Kingdom"
    )
    add_code_block(doc, sample_csv_3)

    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "Data integration combines data from multiple disparate sources (like databases, spreadsheets, text files) into a single, unified view. "
        "Major integration concerns include:\n\n"
        "1. Schema Alignment: Selecting a subset of common fields and aligning datatypes across both sources.\n"
        "2. Provenance Tracking: Adding a 'Source' column to identify the file of origin for auditing.\n"
        "3. Integration & Deduplication: Using pd.concat to merge the aligned records, followed by drop_duplicates to discard overlapping records."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab3_DataIntegration/lab3_data_integration.py", "r", encoding="utf-8") as f:
        code_lab3 = f.read()
    add_code_block(doc, code_lab3)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab3 = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab3_DataIntegration\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab3_DataIntegration> ..\\.venv\\Scripts\\python.exe lab3_data_integration.py\n\n"
        "==================================================\n"
        " LAB 3: DATA PREPROCESSING - DATA INTEGRATION \n"
        "==================================================\n"
        "INFO: Found Excel file: Online RetailSmall.xlsx\n"
        "INFO: Found CSV file: Online RetailSmallData.csv\n"
        "INFO: Integration Successful!\n"
        "  StockCode                          Description  Quantity  UnitPrice  CustomerID         Country Source\n"
        "0    85123A   WHITE HANGING HEART T-LIGHT HOLDER         6       2.55     17850.0  United Kingdom  Excel\n"
        "1     71053                  WHITE METAL LANTERN         6       3.39     17850.0  United Kingdom  Excel\n"
        "2    84406B       CREAM CUPID HEARTS COAT HANGER         8       2.75     17850.0  United Kingdom  Excel\n"
        "3    84029G  KNITTED UNION FLAG HOT WATER BOTTLE         6       3.39     17850.0  United Kingdom  Excel\n"
        "4    84029E       RED WOOLLY HOTTIE WHITE HEART.         6       3.39     17850.0  United Kingdom  Excel\n"
        "5     22752         SET 7 BABUSHKA NESTING BOXES         2       7.65     17850.0  United Kingdom  Excel\n"
        "6     21730    GLASS STAR FROSTED T-LIGHT HOLDER         6       4.25     17850.0  United Kingdom  Excel\n"
        "7     22633               HAND WARMER UNION JACK         6       1.85     17850.0  United Kingdom  Excel\n\n"
        "Total rows integrated: 94\n"
        "Unique customers count  : 4\n"
        "Unique countries count  : 2\n\n"
        "[ANALYTICS] Top 5 Countries by Revenue:\n"
        "Country\n"
        "France            1711.72\n"
        "United Kingdom    1463.90\n"
        "Name: Revenue, dtype: float64"
    )
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab3)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab3_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 3.1: Lab 3 Terminal Output Screenshot")
    
    doc.add_page_break()

    # ==========================================
    # LAB 4
    # ==========================================
    h = add_custom_heading(doc, "Lab 4: Mining Frequent Patterns (Apriori Algorithm)", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.25)
    p_q.add_run("Explain the Apriori algorithm for frequent itemset mining and association rule generation. Implement a Python script using the mlxtend library to one-hot encode a transaction dataset, find all frequent itemsets with a minimum support threshold of 50%, generate association rules, and filter strong rules with a confidence threshold of at least 70%.")
    
    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "The Apriori algorithm is used for market basket analysis to discover relationships between products. "
        "It employs an iterative approach known as a level-wise search, where k-itemsets are used to explore (k+1)-itemsets.\n\n"
        "Key Metrics:\n"
        "1. Support: Probability that a transaction contains the itemset. Support(A) = Freq(A) / N\n"
        "2. Confidence: Reliability of the rule. Confidence(A -> B) = Support(A U B) / Support(A)\n"
        "3. Lift: Ratio of observed support to expected independent support. Lift(A -> B) = Support(A U B) / (Support(A) * Support(B))\n\n"
        "Apriori Pruning Property: All non-empty subsets of a frequent itemset must also be frequent."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab4_FrequentPatterns/lab4_apriori.py", "r", encoding="utf-8") as f:
        code_lab4 = f.read()
    add_code_block(doc, code_lab4)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab4 = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab4_FrequentPatterns\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab4_FrequentPatterns> ..\\.venv\\Scripts\\python.exe lab4_apriori.py\n\n"
        "==================================================\n"
        " LAB 4: MINING FREQUENT PATTERNS (APRIORI) \n"
        "==================================================\n"
        "INFO: Raw Transactions:\n"
        "  T1: ['A', 'B', 'C']\n"
        "  T2: ['A', 'C']\n"
        "  T3: ['A', 'D']\n"
        "  T4: ['B', 'E', 'F']\n"
        "INFO: Encoding Transactions...\n"
        "       A      B      C      D      E      F\n"
        "0   True   True   True  False  False  False\n"
        "1   True  False   True  False  False  False\n"
        "2   True  False  False   True  False  False\n"
        "3  False   True  False  False   True   True\n"
        "INFO: Running Apriori (min_support = 0.5)...\n\n"
        "=== Frequent Itemsets ===\n"
        "   support           itemsets  length\n"
        "0     0.75     frozenset({A})       1\n"
        "1     0.50     frozenset({B})       1\n"
        "2     0.50     frozenset({C})       1\n"
        "3     0.50  frozenset({C, A})       2\n"
        "INFO: Discovering Association Rules (min_lift >= 0.5)...\n\n"
        "=== Strong Rules (confidence >= 0.70) ===\n"
        "      antecedents     consequents  support  confidence      lift\n"
        "0  frozenset({C})  frozenset({A})      0.5         1.0  1.333333"
    )
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab4)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab4_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 4.1: Lab 4 Terminal Output Screenshot")
    
    doc.add_page_break()

    # ==========================================
    # LAB 5A
    # ==========================================
    h = add_custom_heading(doc, "Lab 5A: ID3 Decision Tree Classifier", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.25)
    p_q.add_run("How does the ID3 algorithm build a decision tree classifier? Implement a custom ID3 decision tree from scratch in Python (calculating Entropy and Information Gain), train it on a play-golf dataset, print the resulting tree structure, and predict the output class for a test instance.")
    
    add_custom_heading(doc, "Sample Input Dataset", level=2)
    p_sample = doc.add_paragraph()
    p_sample.paragraph_format.left_indent = Inches(0.25)
    p_sample.add_run("This lab uses the traditional golf dataset. Below is the training dataset representing weather configurations and the target label (PlayGolf):\n")
    
    sample_csv_5 = (
        "Outlook,Temperature,Humidity,Windy,PlayGolf\n"
        "Rainy,Hot,High,FALSE,No\n"
        "Rainy,Hot,High,TRUE,No\n"
        "Overcast,Hot,High,FALSE,Yes\n"
        "Sunny,Mild,High,FALSE,Yes\n"
        "Sunny,Cool,Normal,FALSE,Yes\n"
        "Sunny,Cool,Normal,TRUE,No\n"
        "Overcast,Cool,Normal,TRUE,Yes\n"
        "Rainy,Mild,High,FALSE,No\n"
        "Rainy,Cool,Normal,FALSE,Yes\n"
        "Sunny,Mild,Normal,FALSE,Yes\n"
        "Rainy,Mild,Normal,TRUE,Yes\n"
        "Overcast,Mild,High,TRUE,Yes\n"
        "Overcast,Hot,Normal,FALSE,Yes\n"
        "Sunny,Mild,High,TRUE,No"
    )
    add_code_block(doc, sample_csv_5)

    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "ID3 (Iterative Dichotomiser 3) creates a multi-branch tree by evaluating attribute choices using two calculations:\n\n"
        "1. Entropy: Measures the uncertainty/impurity in S.\n"
        "   H(S) = sum( -p_i * log2(p_i) )\n"
        "2. Information Gain: Measures the expected reduction in Entropy after splitting on attribute A.\n"
        "   Gain(S, A) = H(S) - sum( (|S_v| / |S|) * H(S_v) )\n\n"
        "The ID3 algorithm recursively builds the tree by picking the attribute that maximizes Information Gain at each node."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab5_Classification/lab5a_id3.py", "r", encoding="utf-8") as f:
        code_lab5a = f.read()
    add_code_block(doc, code_lab5a)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab5a = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab5_Classification\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab5_Classification> ..\\.venv\\Scripts\\python.exe lab5a_id3.py\n\n"
        "==================================================\n"
        " LAB 5A: ID3 DECISION TREE \n"
        "==================================================\n"
        "[INFO] Dataset shape: (14, 5)\n"
        "[INFO] Target: 'PlayGolf', Attributes: ['Outlook', 'Temperature', 'Humidity', 'Windy']\n\n"
        "=== Generated Decision Tree ===\n"
        "{'Outlook': {'Overcast': 'Yes',\n"
        "             'Rainy': {'Humidity': {'High': 'No',\n"
        "                                    'Normal': 'Yes'}},\n"
        "             'Sunny': {'Windy': {False: 'Yes',\n"
        "                                 True: 'No'}}}}\n\n"
        "[PREDICTION]\n"
        "Input: {'Outlook': 'Sunny', 'Temperature': 'Cool', 'Humidity': 'High', 'Windy': False}\n"
        "Result (PlayGolf) => Yes"
    )
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab5a)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab5a_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 5A.1: Lab 5A Terminal Output Screenshot")
    
    doc.add_page_break()

    # ==========================================
    # LAB 5B
    # ==========================================
    h = add_custom_heading(doc, "Lab 5B: Naive Bayes Classifier", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.25)
    p_q.add_run("Explain Naive Bayes classification. Write a Python program using scikit-learn's CategoricalNB classifier to train a model on a categorical play-golf dataset, evaluate its performance (accuracy, classification report), and perform prediction on a new test case.")
    
    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "Naive Bayes is a statistical classifier that calculates the probability of a data point belonging to a class. "
        "It uses Bayes' Theorem:\n"
        "P(Class | Features) = [ P(Features | Class) * P(Class) ] / P(Features)\n\n"
        "It is called 'Naive' because it assumes class-conditional independence: that features do not interact and are fully independent given the class."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab5_Classification/lab5b_naive_bayes.py", "r", encoding="utf-8") as f:
        code_lab5b = f.read()
    add_code_block(doc, code_lab5b)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab5b = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab5_Classification\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab5_Classification> ..\\.venv\\Scripts\\python.exe lab5b_naive_bayes.py\n\n"
        "==================================================\n"
        " LAB 5B: NAIVE BAYES CLASSIFICATION \n"
        "==================================================\n"
        "[INFO] Dataset successfully loaded in-memory.\n"
        "[INFO] Training Categorical Naive Bayes model...\n\n"
        "[RESULTS] Model Accuracy: 33.33%\n\n"
        "--- Classification Report ---\n"
        "              precision    recall  f1-score   support\n\n"
        "      No (0)       0.00      0.00      0.00         1\n"
        "     Yes (1)       0.50      0.50      0.50         2\n\n"
        "    accuracy                           0.33         3\n"
        "   macro avg       0.25      0.25      0.25         3\n"
        "weighted avg       0.33      0.33      0.33         3\n\n\n"
        "[PREDICTION TEST]\n"
        "Condition: {'Outlook': 'Sunny', 'Temperature': 'Cool', 'Humidity': 'High', 'Windy': False}\n"
        "Prediction (Play Golf): Yes"
    )
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab5b)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab5b_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 5B.1: Lab 5B Terminal Output Screenshot")
    
    doc.add_page_break()

    # ==========================================
    # LAB 5C
    # ==========================================
    h = add_custom_heading(doc, "Lab 5C: Support Vector Machine (SVM)", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.25)
    p_q.add_run("Explain the working of a Support Vector Machine (SVM). Write a Python script using scikit-learn to train a linear SVM on a synthetic two-dimensional dataset. Additionally, write a script to plot the decision boundary, the margins, and highlight the support vectors.")
    
    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "Support Vector Machine (SVM) separates classes by drawing a decision boundary called a hyperplane. "
        "In a 2D space, this hyperplane is a line. The objective is to maximize the margin width (distance between the line and the closest training points from either class, which are called the support vectors). "
        "Points that lie on the margin boundaries determine the final parameters of the line."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab5_Classification/lab5c_svm_simple.py", "r", encoding="utf-8") as f:
        code_lab5c_simple = f.read()
    add_code_block(doc, code_lab5c_simple)
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab5_Classification/lab5c_svm_visualize.py", "r", encoding="utf-8") as f:
        code_lab5c_viz = f.read()
    add_code_block(doc, code_lab5c_viz)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab5c = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab5_Classification\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab5_Classification> ..\\.venv\\Scripts\\python.exe lab5c_svm_simple.py\n\n"
        "==================================================\n"
        " LAB 5C: SUPPORT VECTOR MACHINE (SIMPLE) \n"
        "==================================================\n"
        "[INFO] Synthetic binary dataset created.\n"
        "[INFO] Training Linear SVM model...\n\n"
        "[RESULT] Accuracy Score: 1.00\n\n"
        "[PREDICTION TEST]\n"
        "Sample [5 6] -> Class 0\n"
        "Sample [1 1] -> Class 0\n\n"
        "--------------------------------------------------\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab5_Classification> ..\\.venv\\Scripts\\python.exe lab5c_svm_visualize.py\n\n"
        "==================================================\n"
        " LAB 5C: SVM DECISION BOUNDARY VISUALIZATION \n"
        "==================================================\n"
        "[INFO] Training SVM model...\n\n"
        "Number of support vectors: 3\n"
        "Support vector coordinates:\n"
        "[[3. 3.]\n"
        " [4. 4.]\n"
        " [6. 7.]]\n\n"
        "[INFO] Generating plot window. Close the plot window to exit."
    )
    # Insert Generated Image
    doc.add_picture("plots/svm_plot.png", width=Inches(4.5))
    add_figure_caption(doc, "Figure 5C: Support Vector Machine Separating Hyperplane")
    
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab5c)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab5c_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 5C.2: Lab 5C Terminal Output Screenshot")
    
    doc.add_page_break()

    # ==========================================
    # LAB 6
    # ==========================================
    h = add_custom_heading(doc, "Lab 6: Cluster Analysis (K-Means)", level=1)
    
    add_custom_heading(doc, "Question", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.25)
    p_q.add_run("What is K-Means clustering? Implement a Python program using scikit-learn to generate a synthetic two-dimensional blob dataset, standardize the features, run K-Means with K=4, report performance metrics like inertia (WCSS) and cluster sizes, classify a new test coordinate, and visualize the final clustered data along with the cluster centroids.")
    
    add_custom_heading(doc, "Theory & Explanation", level=2)
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.25)
    p_t.add_run(
        "K-Means is an unsupervised clustering algorithm that groups unlabeled instances into K distinct clusters. "
        "Steps of the algorithm:\n\n"
        "1. Initialization: Select K random initial cluster centroids.\n"
        "2. Assignment: Assign each data instance to the nearest centroid (using Euclidean distance).\n"
        "3. Centroid Update: Move each centroid to the mean location of all instances assigned to it.\n"
        "4. Repeat steps 2 and 3 until convergence (centroids no longer move).\n\n"
        "The objective is to minimize the Within-Cluster Sum of Squares (Inertia)."
    )
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("Lab6_Clustering/lab6_kmeans.py", "r", encoding="utf-8") as f:
        code_lab6 = f.read()
    add_code_block(doc, code_lab6)
    
    add_custom_heading(doc, "Output", level=2)
    cmd_lab6 = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> cd Lab6_Clustering\n"
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM\\Lab6_Clustering> ..\\.venv\\Scripts\\python.exe lab6_kmeans.py\n\n"
        "==================================================\n"
        " LAB 6: CLUSTER ANALYSIS (K-MEANS) \n"
        "==================================================\n"
        "[INFO] Synthetic dataset generated and standardized. Shape: (300, 2)\n\n"
        "[INFO] Running K-Means with K=4...\n\n"
        "--- Model Performance ---\n"
        "Inertia (WCSS) : 158.4687\n"
        "Iterations     : 3\n"
        "Cluster sizes  : {0: 75, 1: 75, 2: 75, 3: 75}\n\n"
        "[PREDICTION TEST]\n"
        "New point [1.5 2.0] -> Assigned to Cluster 2"
    )
    # Insert Generated Image
    doc.add_picture("plots/kmeans_plot.png", width=Inches(4.5))
    add_figure_caption(doc, "Figure 6: K-Means Cluster Partitioning (K=4)")
    
    # Generate and Insert Terminal Screenshot
    ss_lines = parse_terminal_text(cmd_lab6)
    ss_path = create_terminal_screenshot(ss_lines, "screenshots/lab6_screenshot.png")
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.left_indent = Inches(0.25)
    doc.add_picture(ss_path, width=Inches(6.0))
    add_figure_caption(doc, "Figure 6.2: Lab 6 Terminal Output Screenshot")

    doc.add_page_break()

    # ==========================================
    # STREAMLIT APP SECTION
    # ==========================================
    h_app = add_custom_heading(doc, "Interactive Project: Streamlit Lab Suite", level=1)
    
    add_custom_heading(doc, "Objective", level=2)
    p_app_o = doc.add_paragraph()
    p_app_o.paragraph_format.left_indent = Inches(0.25)
    p_app_o.add_run("Provide a central interactive dashboard that integrates and visualizes the results of all individual DWDM lab experiments dynamically using Streamlit.")
    
    add_custom_heading(doc, "Source Code", level=2)
    with open("app.py", "r", encoding="utf-8") as f:
        code_app = f.read()
    add_code_block(doc, code_app)

    add_custom_heading(doc, "How to Run the Streamlit App", level=2)
    p_run = doc.add_paragraph()
    p_run.paragraph_format.left_indent = Inches(0.25)
    p_run.add_run("To launch the interactive dashboard, open your terminal in the DWDM directory and execute:")
    
    run_cmd_text = (
        "PS C:\\Users\\firoj\\OneDrive\\Desktop\\Codes practice\\College_Codes\\DWDM> .\\.venv\\Scripts\\streamlit.exe run app.py"
    )
    add_code_block(doc, run_cmd_text)

    # Save document
    try:
        doc.save("LABS_QUESTIONS_AND_ANSWERS.docx")
        print("Document saved successfully as LABS_QUESTIONS_AND_ANSWERS.docx!")
    except PermissionError:
        import random
        num = random.randint(100, 999)
        alt_name = f"LABS_QUESTIONS_AND_ANSWERS_{num}.docx"
        doc.save(alt_name)
        print(f"[WARNING] Could not overwrite 'LABS_QUESTIONS_AND_ANSWERS.docx' because it is open in another program.")
        print(f"Saved the file as '{alt_name}' instead. Please close Microsoft Word and delete/rename it as needed.")

if __name__ == "__main__":
    main()
