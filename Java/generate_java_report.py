import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import subprocess

def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    leading_spaces = len(text) - len(text.lstrip(' '))
    p.paragraph_format.left_indent = Inches(0.25 + (leading_spaces * 0.05))
    
    run = p.add_run(text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

def add_output_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def clean_code_line(text):
    text = text.replace('\xa0', ' ').replace('\u200b', '')
    dot_idx = text.find('.')
    if dot_idx != -1 and dot_idx < 8:
        prefix = text[:dot_idx].strip()
        cleaned_prefix = ''.join([c for c in prefix if c.isdigit()])
        if cleaned_prefix:
            return text[dot_idx + 1:]
    return text

def find_file_and_run(code_text):
    # Find class name
    match = re.search(r'public\s+class\s+([a-zA-Z0-9_]+)', code_text)
    if not match:
        match = re.search(r'class\s+([a-zA-Z0-9_]+)', code_text)
    if not match:
        return "No class declaration found in code block."
        
    class_name = match.group(1)
    
    # Find package name
    pkg_match = re.search(r'package\s+([a-zA-Z0-9_\.]+);', code_text)
    package_prefix = ""
    if pkg_match:
        package_prefix = pkg_match.group(1) + "."
        
    full_class_name = package_prefix + class_name
    
    # Determine if it's a GUI/Web/Network task
    is_gui_or_web = False
    for indicator in ["javax.swing", "javafx", "javax.servlet", "java.rmi", "ServerSocket", "Socket", "DatagramSocket", "DriverManager.getConnection"]:
        if indicator in code_text:
            is_gui_or_web = True
            break
            
    if is_gui_or_web:
        java_root = "C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes/Java"
        rel_folder = "Java/Lab"
        for dirpath, _, filenames in os.walk(java_root):
            for f in filenames:
                if f == f"{class_name}.java":
                    rel_folder = os.path.relpath(dirpath, "C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes")
                    break
        return f"[GUI / Web / Network Task: Please run this task manually in your environment by navigating to '{rel_folder.replace(chr(92), '/')}' and executing the program or deploy.ps1 script.]"

    # Search for the source file
    java_root = "C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes/Java"
    found_path = None
    for dirpath, _, filenames in os.walk(java_root):
        if "class" in dirpath or "temp_dont_push" in dirpath:
            continue
        for filename in filenames:
            if filename == f"{class_name}.java":
                found_path = os.path.join(dirpath, filename)
                break
        if found_path:
            break
            
    if not found_path:
        return f"[Source file {class_name}.java not found in workspace. Run manually.]"
        
    # Compile
    out_dir = "C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes/Java/class"
    os.makedirs(out_dir, exist_ok=True)
    compile_cmd = ["javac", "-d", out_dir, found_path]
    
    try:
        subprocess.run(compile_cmd, capture_output=True, text=True, check=True)
    except subprocess.CalledProcessError as e:
        return f"Compilation Error:\n{e.stderr}"
        
    # Run
    run_cmd = ["java", "-cp", out_dir, full_class_name]
    try:
        res = subprocess.run(run_cmd, capture_output=True, text=True, timeout=3)
        output_data = res.stdout
        if res.stderr:
            output_data += "\n" + res.stderr
        return output_data.strip()
    except subprocess.TimeoutExpired:
        return "[Execution timed out (infinite loop or waiting for input).]"
    except Exception as e:
        return f"[Execution Error: {str(e)}]"

def main():
    print("Reading template Java_Labs.docx...")
    try:
        orig_doc = Document("C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes/Java/temp_dont_push/Java_Labs.docx")
    except Exception as e:
        print(f"Error loading Java_Labs.docx: {e}")
        return

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0, 0, 0)

    in_code = False
    code_lines = []

    print("Processing document sections...")
    for para in orig_doc.paragraphs:
        text = para.text.rstrip()
        if not text:
            doc.add_paragraph("")
            continue

        if text.lower().startswith("code"):
            in_code = True
            code_lines = []
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        if text.lower().startswith("output"):
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            
            # Print the actual run output dynamically
            print("Running code block to capture stdout...")
            code_text = "\n".join(code_lines)
            captured_output = find_file_and_run(code_text)
            
            add_output_paragraph(doc, captured_output)
            continue

        if in_code:
            cleaned = clean_code_line(text)
            code_lines.append(cleaned)
            add_code_paragraph(doc, cleaned)
        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            if "write a program" in text.lower() or "create a" in text.lower() or "develop a" in text.lower() or "create base class" in text.lower():
                run.font.bold = True

    # Save output
    try:
        doc.save("C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes/Java/Java_Lab_Report.docx")
        print("Beautiful document saved successfully as 'Java_Lab_Report.docx'!")
    except PermissionError:
        import random
        num = random.randint(100, 999)
        alt_name = f"C:/Users/firoj/OneDrive/Desktop/Codes practice/College_Codes/Java/Java_Lab_Report_{num}.docx"
        doc.save(alt_name)
        print(f"[WARNING] Could not overwrite 'Java_Lab_Report.docx' because it is open in another program.")
        print(f"Saved the file as '{alt_name}' instead. Please close Microsoft Word.")

if __name__ == "__main__":
    main()
